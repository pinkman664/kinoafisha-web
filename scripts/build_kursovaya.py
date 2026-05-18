"""
Генерация курсового проекта «Киноафиша» по СТП БГТУ 001-2019.
Создаёт единый файл /app/Курсовая.docx, в который далее будут
дописываться разделы.
"""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_page_setup(doc: Document) -> None:
    """Поля по СТП БГТУ: правое 10 мм, верхнее 20 мм, левое 23 мм,
    нижнее 15 мм. Номер страницы — 10 мм от верхнего края."""
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(23)
        section.right_margin = Mm(10)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


def add_page_number_header(doc: Document) -> None:
    """Номер страницы в правом верхнем углу через автополе PAGE."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_default_style(doc: Document) -> None:
    """Times New Roman 14 пт, одинарный межстрочный интервал."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_heading_special(doc: Document, text: str) -> None:
    """Специальный заголовок (Введение, Заключение, Содержание и т.п.):
    строчные кроме первой прописной, полужирный, симметрично тексту,
    без точки, интервал 18 пт после заголовка."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.first_line_indent = Mm(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_body_paragraph(doc: Document, text: str) -> None:
    """Обычный абзац: красная строка 12,5 мм, выравнивание по ширине."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(12.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def main() -> None:
    doc = Document()

    set_page_setup(doc)
    set_default_style(doc)
    add_page_number_header(doc)

    # ────────────────────────────────────────────────────────────────────
    # ВВЕДЕНИЕ
    # ────────────────────────────────────────────────────────────────────
    add_heading_special(doc, "Введение")

    add_body_paragraph(
        doc,
        "Посещение кинотеатров остаётся одной из наиболее массовых форм "
        "досуга, а сам кинопрокат — крупной отраслью индустрии развлечений. "
        "За последние годы существенно изменился сам способ взаимодействия "
        "зрителя с кинотеатром: продажа билетов через кассу постепенно "
        "вытесняется онлайн-сервисами, которые позволяют ознакомиться с "
        "афишей, выбрать конкретный сеанс, забронировать места в зале и "
        "оплатить заказ непосредственно с компьютера или мобильного "
        "устройства. Подобные сервисы становятся обязательным элементом "
        "работы современных кинотеатров и одновременно удобным инструментом "
        "для зрителей, что определяет актуальность разработки полнофункцио"
        "нального web-приложения подобного класса."
    )

    add_body_paragraph(
        doc,
        "Разрабатываемое web-приложение «Киноафиша» предназначено для "
        "автоматизации деятельности кинотеатра в части представления "
        "репертуара и продажи билетов на сеансы. Приложение объединяет в "
        "единой среде функции просмотра афиши, изучения детальной информации "
        "о фильмах с возможностью воспроизведения трейлеров, поиска и "
        "фильтрации картин по жанрам и зрительскому рейтингу, выбора "
        "конкретного места в зале и оплаты заказа через интернет-эквайринг. "
        "После успешной оплаты пользователь получает электронный билет с "
        "QR-кодом, который сохраняется в личном кабинете и предъявляется на "
        "входе в зал. Помимо функций, ориентированных на зрителя, в "
        "приложении предусмотрена административная подсистема, позволяющая "
        "управлять списком фильмов, классификаторами жанров и кинотеатров, "
        "схемами залов, расписанием сеансов, пользователями и публикуемыми "
        "отзывами, а также получать сводную статистику по продажам."
    )

    add_body_paragraph(
        doc,
        "Доступ к функциональным возможностям приложения разграничивается "
        "при помощи ролевой модели: к роли «Гость» относятся незарегистри"
        "рованные посетители, имеющие возможность только просматривать "
        "информацию о фильмах и сеансах; роль «Клиент» получают зарегистри"
        "рованные пользователи, которым доступно бронирование и оплата "
        "билетов, отмена ранее купленных билетов, публикация отзывов и "
        "выставление оценок; роль «Администратор» обладает полным набором "
        "прав по управлению содержимым приложения и его пользователями."
    )

    add_body_paragraph(
        doc,
        "Целью курсового проекта является разработка web-приложения "
        "«Киноафиша», обеспечивающего полный цикл взаимодействия пользователя "
        "с кинотеатром от выбора фильма и сеанса до получения электронного "
        "билета, а также предоставляющего администратору средства управления "
        "репертуаром и анализа результатов продаж. Для достижения "
        "поставленной цели необходимо провести анализ предметной области и "
        "обзор существующих аналогов, что рассматривается в главе 1, "
        "спроектировать архитектуру приложения, разработать модель базы "
        "данных и выполнить программную реализацию серверной и клиентской "
        "частей системы, представленные в главе 2, а также выполнить "
        "тестирование разработанного приложения и составить руководство "
        "пользователя, рассмотренные в главе 3."
    )

    add_body_paragraph(
        doc,
        "Целевой аудиторией приложения являются зрители, выбирающие фильмы "
        "и приобретающие билеты онлайн, а также сотрудники и администраторы "
        "сети кинотеатров, заинтересованные в централизованном инструменте "
        "управления репертуаром, расписанием сеансов и анализе спроса."
    )

    add_body_paragraph(
        doc,
        "В качестве программной платформы используются Node.js 22, "
        "TypeScript 5, web-фреймворк Express 5 и объектно-реляционное "
        "отображение TypeORM 0.3 для серверной части; библиотека React 18.3 "
        "и сборщик Vite 5 для клиентской части; система управления базами "
        "данных Oracle Database 21c; платёжный сервис ЮKassa для приёма "
        "электронных платежей; средства контейнеризации Docker 27 и Docker "
        "Compose, а также прокси-сервер Nginx 1.27 для маршрутизации "
        "запросов и раздачи статических ресурсов."
    )

    output_path = "/app/Курсовая.docx"
    doc.save(output_path)
    print(f"Документ сохранён: {output_path}")


if __name__ == "__main__":
    main()
