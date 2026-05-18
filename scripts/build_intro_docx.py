"""
Скрипт генерации Введения курсового проекта по СТП БГТУ 001-2019.
Создаёт файл /app/Введение.docx.
"""
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_page_setup(doc: Document) -> None:
    """Поля: правое 10 мм, верхнее 20 мм, левое 23 мм, нижнее 15 мм."""
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(23)
        section.right_margin = Mm(10)
        # Номер страницы — в правом верхнем углу, 10 ± 2 мм от верхней границы
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


def add_page_number_header(doc: Document) -> None:
    """Добавляет номер страницы в правый верхний угол.
    Титульный лист обычно делается отдельно, на нём номер не ставится.
    Поэтому здесь делаем different_first_page = True, и оставляем
    первую страницу без номера (под Введение это не критично — обычно
    Введение начинается с нового листа и нумеруется)."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    # PAGE field
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_default_style(doc: Document) -> None:
    """Times New Roman 14 пт, одинарный межстрочный интервал."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    # Для корректной работы кириллицы
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_heading_special(doc: Document, text: str) -> None:
    """Специальный заголовок (Введение, Заключение и т.п.):
    - строчные буквы, кроме первой прописной;
    - полужирный;
    - симметрично тексту (по центру);
    - без точки;
    - интервал 18 пт после заголовка.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.first_line_indent = Mm(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_body_paragraph(doc: Document, text: str) -> None:
    """Обычный абзац: красная строка 12,5 мм, выравнивание по ширине."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(12.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_list_item(doc: Document, text: str) -> None:
    """Пункт перечисления: с абзацного отступа, начинается с дефиса
    и строчной буквы; в конце — точка с запятой (последний — точка).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(12.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def main() -> None:
    doc = Document()

    set_page_setup(doc)
    set_default_style(doc)
    add_page_number_header(doc)

    add_heading_special(doc, "Введение")

    add_body_paragraph(
        doc,
        "В современном мире индустрия кинопроката переживает значительные "
        "изменения, связанные с переходом большинства бизнес-процессов в "
        "цифровую среду. Зрители всё реже обращаются непосредственно в кассы "
        "кинотеатров, предпочитая выбирать сеансы, бронировать места и "
        "оплачивать билеты в режиме онлайн с использованием персональных "
        "компьютеров и мобильных устройств. Кинотеатры, в свою очередь, "
        "заинтересованы в инструментах, позволяющих автоматизировать продажу "
        "билетов, управлять расписанием сеансов и анализировать спрос на "
        "различные фильмы. В этих условиях web-приложение, объединяющее "
        "интересы зрителей и администраторов кинотеатров на единой платформе, "
        "становится востребованным программным продуктом."
    )

    add_body_paragraph(
        doc,
        "Актуальность темы курсового проектирования обусловлена постоянно "
        "растущим спросом на удобные и надёжные сервисы онлайн-бронирования "
        "билетов, а также необходимостью практического освоения современных "
        "подходов к разработке полнофункциональных web-приложений с "
        "разделением ролей пользователей, интеграцией с платёжными системами "
        "и хранением данных в реляционных базах данных промышленного уровня."
    )

    add_body_paragraph(
        doc,
        "Целью курсового проекта является разработка web-приложения "
        "«Киноафиша», обеспечивающего полный цикл взаимодействия пользователя "
        "с кинотеатром — от поиска и просмотра информации о фильмах до "
        "приобретения электронного билета с QR-кодом."
    )

    add_body_paragraph(
        doc,
        "Для достижения поставленной цели в работе решаются следующие задачи:"
    )

    items = [
        "– провести анализ предметной области и обзор существующих "
        "аналогичных решений, представленных на рынке;",

        "– определить функциональные и нефункциональные требования к "
        "разрабатываемому приложению с учётом трёх ролей пользователей: "
        "гостя, клиента и администратора;",

        "– спроектировать архитектуру приложения, обеспечивающую разделение "
        "представления, бизнес-логики и хранилища данных, а также разработать "
        "модель базы данных;",

        "– реализовать серверную часть приложения с использованием платформы "
        "Node.js и языка TypeScript, организовав взаимодействие с клиентом на "
        "базе архитектурного стиля REST;",

        "– реализовать клиентскую часть в виде одностраничного приложения "
        "(SPA) с применением библиотеки React;",

        "– интегрировать приложение с платёжной системой для приёма онлайн-"
        "платежей и автоматической генерации электронных билетов;",

        "– выполнить контейнеризацию приложения средствами Docker и Docker "
        "Compose, организовать маршрутизацию запросов через прокси-сервер "
        "Nginx;",

        "– провести тестирование разработанного web-приложения и составить "
        "руководство пользователя.",
    ]
    for item in items:
        add_list_item(doc, item)

    add_body_paragraph(
        doc,
        "Объектом разработки является процесс взаимодействия зрителей и "
        "администраторов кинотеатра в рамках электронной системы бронирования "
        "и продажи билетов на киносеансы."
    )

    add_body_paragraph(
        doc,
        "Предметом разработки выступают программные и алгоритмические "
        "средства, обеспечивающие реализацию указанного взаимодействия в "
        "форме web-приложения с асинхронным пользовательским интерфейсом."
    )

    add_body_paragraph(
        doc,
        "В качестве основного языка разработки используется TypeScript. "
        "Серверная часть построена на платформе Node.js с применением web-"
        "фреймворка Express и объектно-реляционного отображения TypeORM. "
        "Клиентская часть реализована на библиотеке React с использованием "
        "сборщика Vite. В качестве системы управления базами данных "
        "применяется Oracle Database. Для обработки платежей используется "
        "сервис ЮKassa. Развёртывание приложения осуществляется в Docker-"
        "контейнерах с применением Docker Compose, маршрутизация запросов и "
        "раздача статических ресурсов выполняется прокси-сервером Nginx с "
        "поддержкой протокола HTTPS."
    )

    add_body_paragraph(
        doc,
        "Пояснительная записка состоит из введения, пяти основных разделов, "
        "заключения, списка использованных источников и приложений. В первом "
        "разделе формулируется постановка задачи и приводится обзор "
        "аналогичных программных решений. Во втором разделе рассматриваются "
        "вопросы проектирования web-приложения, включая выбор архитектуры, "
        "проектирование базы данных и интерфейса пользователя. Третий раздел "
        "посвящён реализации программного средства. В четвёртом разделе "
        "описаны методика и результаты тестирования. В пятом разделе "
        "приведено руководство пользователя."
    )

    output_path = "/app/Введение.docx"
    doc.save(output_path)
    print(f"Документ сохранён: {output_path}")


if __name__ == "__main__":
    main()
